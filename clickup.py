import os
import re
import time
import json
import pandas as pd
import requests
import google.generativeai as genai
from docx import Document
from dotenv import load_dotenv

#CARGAR CREDENCIALES
load_dotenv()
CLICKUP_TOKEN = os.getenv('CLICKUP_TOKEN')
GEMINI_TOKEN = os.getenv('GEMINI_TOKEN')

genai.configure(api_key=GEMINI_TOKEN)

#CONFIGURACIÓN PERSONALIZADA
LIST_IDS = ['LISTA DE ID CLICKUP'] 
PREFIJOS_VALIDOS = ('SI LA TAREA DE CLICKUP TIENE PREFIJOS')

def leer_reglas_qa():
    try:
        with open('prompt-test-cases (1).md', 'r', encoding='utf-8') as archivo:
            return archivo.read()
    except FileNotFoundError:
        print(" Error: No se encontró el archivo de reglas.")
        return ""

def limpiar_nombre_carpeta(nombre):
    """Limpia caracteres inválidos para crear carpetas en Windows/Mac"""
    return re.sub(r'[\\/*?:"<>|]', "", nombre).strip()

def generar_casos_con_gemini(tarea, reglas_qa):
    print(f"\n🧠 Pensando casos para: {tarea['nombre']}...")
    
    descripcion_segura = tarea['descripcion'][:6000] 
    
    modelo = genai.GenerativeModel(
        model_name='gemini-3.5-flash', 
        system_instruction=reglas_qa 
    )
    

    prompt_usuario = f"""
    Genera los casos de prueba para esta tarea de BeOS basándote en mis reglas.
    
    TÍTULO DE LA TAREA: {tarea['nombre']}
    
    DESCRIPCIÓN Y CRITERIOS: 
    {descripcion_segura}

    🎯 REGLA DE CANTIDAD DE CASOS (¡ESTRICTO!):
    - Analiza la complejidad de la tarea descrita arriba.
    - Si la tarea es MUY MÍNIMA, genera de 3 a 5 casos de prueba.
    - Si la tarea es de complejidad MEDIA o ALTA, ESTÁS OBLIGADO a generar un MÍNIMO DE 8 CASOS DE PRUEBA (Flujos negativos, roles, edge cases).

    ⚠️ INSTRUCCIÓN OBLIGATORIA DE FORMATO JSON:
    Tu respuesta debe ser un arreglo JSON con las siguientes llaves EXACTAS:
    [
        {{
            "Title": "TC-XXX-001 | Título del caso",
            "Description": "Descripción detallada del caso.",
            "Preconditions": "Condiciones previas...",
            "Steps": "1. Identificar una admisión. 2. Generar el reporte. 3. Inspeccionar el archivo. -> Las columnas muestran explícitamente el texto.",
            "Expected Result": "Resultado esperado final",
            "Severity": "Normal",
            "Priority": "High",
            "Suite ID": "",
            "Risk Tier": "Tier 1",
            "Automation Status": "Not automated",
            "Execution Environment": "v100"
        }}
    ]
    
    🛠️ REGLAS DE LLENADO CRÍTICAS PARA "STEPS":
    - ¡PROHIBIDO USAR SALTOS DE LÍNEA (\\n)! 
    - Escribe TODOS los pasos de forma horizontal en una sola línea continua, separándolos únicamente con un espacio. 
    - Nunca dejes la prueba en un solo paso. Desglosa en paso 1, 2, 3, etc.
    - Al final del último paso, DEBES colocar una flecha '->' seguida de la validación o resultado visual esperado en ese paso.
    - EJEMPLO EXACTO QUE DEBES IMITAR: "1. Identificar una admisión con marca Vol/Invol en 'Yes' y otra en 'No'. 2. Generar el reporte Admission Tracking Report - Inpatient. 3. Inspeccionar el archivo generado. -> Las columnas del Excel muestran explícitamente el texto 'Vol' e 'Invol' en lugar de los valores booleanos anteriores."
    
    OTRAS REGLAS:
    - Deja "Suite ID" EN BLANCO ("") para que el QA lo llene manualmente después.
    - Severity: Trivial, Minor, Normal, Major, Critical o Blocker.
    - Priority: Low, Medium o High.
    """
    
    print("   Enviando petición a Gemini API")
    try:
        respuesta = modelo.generate_content(
            prompt_usuario,
            request_options={"timeout": 120}, 
            generation_config={"response_mime_type": "application/json"} 
        )
        print("   ¡Casos generados con éxito!")
        return respuesta.text
    except Exception as e:
        print(f"   Falló la generación. Motivo: {e}")
        return None

def crear_documento_word(casos_dict, ruta_word, titulo_tarea):
    """Convierte el JSON de casos en un documento Word bonito"""
    doc = Document()
    doc.add_heading(f"Casos de Prueba: {titulo_tarea}", 0)
    
    for i, caso in enumerate(casos_dict, 1):
        doc.add_heading(caso.get('Title', f'Caso {i}'), level=1)
        
        doc.add_heading(" Descripción:", level=2)
        doc.add_paragraph(caso.get('Description', ''))
        
        doc.add_heading(" Pre-condiciones:", level=2)
        doc.add_paragraph(caso.get('Preconditions', 'Ninguna'))
        
        doc.add_heading(" Pasos:", level=2)
        # Los pasos ahora se insertan tal cual (horizontales)
        doc.add_paragraph(caso.get('Steps', ''))
        
        doc.add_heading(" Resultado Esperado:", level=2)
        doc.add_paragraph(caso.get('Expected Result', ''))
        
        doc.add_heading(" Metadatos:", level=2)
        doc.add_paragraph(f"Severity: {caso.get('Severity', '')} | Priority: {caso.get('Priority', '')} | Risk Tier: {caso.get('Risk Tier', '')}")
        
        doc.add_page_break() 
        
    doc.save(ruta_word)

def analizar_y_generar():
    print("Iniciando extracción de ClickUp...\n")
    tareas_listas_para_ia = []
    tareas_rechazadas = []
    
    for list_id in LIST_IDS:
        url = f"https://api.clickup.com/api/v2/list/{list_id}/task"
        headers = {"Authorization": CLICKUP_TOKEN}
        try:
            response = requests.get(url, headers=headers, timeout=15)
            if response.status_code == 200:
                for task in response.json().get('tasks', []):
                    task_name = task.get('name', '')
                    desc = task.get('description', '')
                    task_url = task.get('url', 'Enlace no disponible')
                    
                    if not task_name.upper().startswith(PREFIJOS_VALIDOS):
                        tareas_rechazadas.append({
                            "nombre": task_name,
                            "razon": f"No tiene un prefijo válido {PREFIJOS_VALIDOS}",
                            "enlace": task_url
                        })
                        continue 
                        
                    if not desc or len(desc.strip()) < 20:
                        tareas_rechazadas.append({
                            "nombre": task_name,
                            "razon": "Descripción vacía o menor a 20 caracteres (Falta de Criterios)",
                            "enlace": task_url
                        })
                        continue
                        
                    tareas_listas_para_ia.append({"nombre": task_name, "descripcion": desc})
        except Exception as e:
            print(f"  ClickUp falló: {e}")
            continue
            
    print(f"\n TAREAS APROBADAS PARA IA: {len(tareas_listas_para_ia)}")
    print("=" * 50)
    
    reglas_qa = leer_reglas_qa()
    if not reglas_qa or len(tareas_listas_para_ia) == 0:
        print("No hay tareas válidas para procesar o falta el archivo de reglas. Terminando...")
    else:
        if not os.path.exists("Entregables_QA"):
            os.makedirs("Entregables_QA")

        for tarea in tareas_listas_para_ia:
            casos_generados_json = generar_casos_con_gemini(tarea, reglas_qa)
            
            if casos_generados_json:
                try:
                    casos_dict = json.loads(casos_generados_json)
                    nombre_seguro = limpiar_nombre_carpeta(tarea['nombre'])
                    
                    ruta_carpeta_tarea = os.path.join("Entregables_QA", nombre_seguro)
                    os.makedirs(ruta_carpeta_tarea, exist_ok=True)
                    
                    # Generar Excel
                    ruta_excel = os.path.join(ruta_carpeta_tarea, f"{nombre_seguro}.xlsx")
                    df = pd.DataFrame(casos_dict)
                    df.to_excel(ruta_excel, index=False)
                    
                    # Generar Word
                    ruta_word = os.path.join(ruta_carpeta_tarea, f"{nombre_seguro}.docx")
                    crear_documento_word(casos_dict, ruta_word, tarea['nombre'])
                    
                    print(f"   ¡ÉXITO! Se crearon Word y Excel en la carpeta: {nombre_seguro}")
                    
                except json.JSONDecodeError:
                    print("  Error procesando el formato de IA.")
            
          
            print(" Esperando 5 segundos para la siguiente tarea...")
            time.sleep(5)
            
    print("\n ¡PROCESO COMPLETADO! Revisa la carpeta 'Entregables_QA'.")

    if tareas_rechazadas:
        print("\n" + "!" * 70)
        print(" REPORTE DE AUDITORÍA: TAREAS RECHAZADAS (NO PROCESADAS)")
        print("!" * 70)
        for rechazada in tareas_rechazadas:
            print(f" Tarea: {rechazada['nombre']}")
            print(f"   Motivo: {rechazada['razon']}")
            print(f"   Enlace: {rechazada['enlace']}\n")
        print("!" * 70)

if __name__ == "__main__":
    analizar_y_generar()